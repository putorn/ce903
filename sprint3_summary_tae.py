"""Creates Sprint 3 academic summary Word document for Tae's group submission."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins (academic standard) ──────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Default font ───────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ── Helpers ────────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p.runs[0].font.name = "Times New Roman"
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p.runs[0].font.name = "Times New Roman"
    return p

def para(text, bold=False, italic=False, indent=False, spacing_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.5)
    p.paragraph_format.space_after   = Pt(3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(80, 80, 80)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D6E4F0")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            if c_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f"[ {text} ]")
    run.font.color.rgb = RGBColor(160, 160, 160)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

# ── Title block ───────────────────────────────────────────────────────────────
title = doc.add_heading("Sprint 3 Technical Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run("Role 2: ASTRA-sim & Simulation Setup")
r.font.name = "Times New Roman"
r.font.size = Pt(13)
r.bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(16)
r = meta.add_run(
    f"CE903 Group Project 3  ·  Managing Network for LLM Training\n"
    f"Tae  ·  {datetime.date.today().strftime('%d %B %Y')}"
)
r.font.name = "Times New Roman"
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ── Overview ──────────────────────────────────────────────────────────────────
h1("1. Overview")

para(
    "This report documents the Sprint 3 deliverables for Role 2 (ASTRA-sim and Simulation Setup). "
    "The sprint objective was to integrate the separate Sprint 2 artefacts into a working simulation "
    "pipeline. Two acceptance tests are addressed: TC-01, which validates the end-to-end "
    "trace-to-simulator pipeline, and TC-02, which validates ring all-reduce scaling behaviour "
    "at GPT-2 124M collective sizes. Both tests have been executed and passed. "
    "Coordination items for other roles have been communicated separately via the group channel."
)

doc.add_paragraph()

# ── TC-01 ─────────────────────────────────────────────────────────────────────
h1("2. TC-01 — Trace-to-Sim Pipeline")

h2("2.1 Objective")
para(
    "TC-01 requires that a single training job runs to completion through the full "
    "trace-to-simulator pipeline with metrics recorded correctly. This is documented as a "
    "Sprint 2 carry-over in the SRS and closes in the current sprint."
)

h2("2.2 Methodology")
para(
    "A representative synthetic Chakra execution trace was constructed to model one "
    "GPT-2 124M training step across four NPUs. The trace encodes four sequential nodes "
    "connected by data dependencies: a forward pass, a backward pass, an AllReduce "
    "collective (512 MB FP32, consistent with the Sprint 3 collective specification), "
    "and an optimiser step. The trace was ingested by the ASTRA-sim 2.0 analytical backend "
    "operating on a Ring topology (50 GB/s, 500 ns latency). Compute timings are "
    "representative estimates derived from published GPT-2 124M profiling data; "
    "real profiler timings will replace these values once the converted trace is delivered "
    "by Role 3."
)

h2("2.3 Results")
para("Simulation completed successfully for all four NPUs with no errors.")
doc.add_paragraph()

add_table(
    headers=["Metric", "Value", "Notes"],
    rows=[
        ["JCT (Wall time)", "95,767,236 cycles", "~95.8 ms total step time"],
        ["GPU time", "80,000,000 cycles (83.5%)", "forward + backward + optimiser"],
        ["Comm time", "15,767,236 cycles (16.5%)", "AllReduce 512 MB FP32"],
        ["Errors", "None", "err.log empty for all NPUs"],
        ["NPU symmetry", "All identical", "ring collective is symmetric"],
    ],
    col_widths=[1.8, 2.5, 2.5]
)
caption("Table 1. TC-01 simulation results — synthetic GPT-2 124M trace, 4 NPUs, Ring topology.")

doc.add_paragraph()
placeholder("TC-01 validate_tc01.py terminal output screenshot")
caption("Figure 1. TC-01 validation output confirming PASS on all four criteria.")

doc.add_paragraph()

h2("2.4 Acceptance Verdict")
para("TC-01 PASSED. The pipeline accepts a Chakra .et trace, simulates compute and "
     "communication phases in the correct dependency order, and records JCT, GPU time, "
     "and Comm time correctly. The 83.5 / 16.5 compute-to-communication split is "
     "consistent with expected GPT-2 training behaviour.")

doc.add_paragraph()

# ── TC-02 ─────────────────────────────────────────────────────────────────────
h1("3. TC-02 — All-Reduce Baseline Sweep")

h2("3.1 Objective")
para(
    "TC-02 requires that simulated ring all-reduce over the baseline analytical topology "
    "produces sensible bandwidth and latency scaling curves with no timeouts. This sweep "
    "constitutes the network-traffic baseline against which all later routing experiments "
    "will be compared."
)

h2("3.2 Methodology")
para(
    "Ring all-reduce microbenchmarks were executed across a three-point NPU sweep "
    "(4, 16, 64) at three message sizes derived from the Sprint 3 collective specification: "
    "1 MB (diagnostic), 256 MB (GPT-2 124M BF16 gradient), and 512 MB (GPT-2 124M FP32 gradient). "
    "All nine simulations used the ASTRA-sim analytical backend with a Ring topology "
    "(50 GB/s, 500 ns latency). Collective synthesis followed the confirmed analytical "
    "approach: no live GPU cluster was required. "
    "TC-02 is evaluated on the two GPT-2 sizes only; the 1 MB run is retained as a "
    "diagnostic to verify algorithm correctness at small scale."
)

h2("3.3 Results")

add_table(
    headers=["Message Size", "NPUs", "Avg Comm (cycles)", "Algo BW (B/cyc)", "Ring Eff."],
    rows=[
        ["256 MB (BF16)", "4",  "398,542",   "673.54", "1.5000"],
        ["256 MB (BF16)", "16", "1,159,300", "231.55", "1.8750"],
        ["256 MB (BF16)", "64", "3,446,957", "77.88",  "1.9688"],
        ["512 MB (FP32)", "4",  "3,630,842", "147.86", "1.5000"],
        ["512 MB (FP32)", "16", "4,539,234", "118.27", "1.8750"],
        ["512 MB (FP32)", "64", "7,554,490", "71.07",  "1.9688"],
    ],
    col_widths=[1.6, 0.7, 1.9, 1.7, 1.0]
)
caption("Table 2. TC-02 simulation results — ring all-reduce at GPT-2 collective sizes.")

doc.add_paragraph()
placeholder("TC-02 analyze_results.py terminal output screenshot")
caption("Figure 2. TC-02 analysis output showing bandwidth plateau and latency scaling.")

doc.add_paragraph()

h2("3.4 Validation")
para("Both GPT-2 message sizes satisfy all TC-02 acceptance criteria:", spacing_after=4)
bullet("Communication time grows monotonically as ring size increases from 4 to 64 NPUs.")
bullet("Algorithm bandwidth does not scale linearly with N — it plateaus as predicted by "
       "the ring all-reduce bandwidth formula (bus BW = algo BW x 2(N-1)/N).")
bullet("Ring efficiency values (1.5000, 1.8750, 1.9688) match theoretical predictions exactly, "
       "confirming correct algorithm implementation.")
bullet("No errors or timeouts were recorded in any of the nine simulation runs.")

doc.add_paragraph()
para(
    "A notable finding at 512 MB FP32 is that bus bandwidth at 4 NPUs (221.80 B/cyc) "
    "and 16 NPUs (221.76 B/cyc) is effectively identical, demonstrating bandwidth "
    "saturation at GPT-2 scale. This result quantifies the point at which adding "
    "ring nodes yields no throughput improvement — a key reference for Sprint 4 "
    "failure and routing experiments."
)

h2("3.5 Acceptance Verdict")
para("TC-02 PASSED on both GPT-2 message sizes (256 MB BF16 and 512 MB FP32).")

doc.add_paragraph()

# ── Repository ────────────────────────────────────────────────────────────────
h1("4. Code and Reproducibility")

para(
    "All simulation scripts, configuration files, and result logs have been committed "
    "to the group repository. The two relevant directories are self-contained and "
    "documented; any team member can reproduce the results by following the instructions "
    "in the respective README files."
)

doc.add_paragraph()

add_table(
    headers=["Directory", "Contents"],
    rows=[
        ["pod_b_traffic/", "TC-02 sweep: workload generation, 9 simulations, analysis script"],
        ["pod_a_pipeline/", "TC-01 pipeline: trace generation, ASTRA-sim run, validation script"],
    ],
    col_widths=[2.2, 4.5]
)
caption("Table 3. Repository directories relevant to this sprint.")

doc.add_paragraph()

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/puthornthongthavee/Desktop/Essex University/project_group/Sprint3_Summary_Tae.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
